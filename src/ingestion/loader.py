from abc import ABC, abstractmethod
from typing import List, Optional, Callable
from pathlib import Path
import re
import logging

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentPreprocessor:
    """Class xử lý preprocessing cho documents"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean text cơ bản
        
        Args:
            text: Raw text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    @staticmethod
    def remove_special_characters(text: str, keep_vietnamese: bool = True) -> str:
        """
        Loại bỏ ký tự đặc biệt
        
        Args:
            text: Input text
            keep_vietnamese: Giữ lại dấu tiếng Việt
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        if keep_vietnamese:
            # Giữ chữ cái (bao gồm tiếng Việt), số, và dấu câu cơ bản
            text = re.sub(r'[^\w\s\.,!?;:\-àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴĐ]', ' ', text)
        else:
            # Chỉ giữ ASCII
            text = re.sub(r'[^\w\s\.,!?;:\-]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    @staticmethod
    def normalize_whitespace(text: str) -> str:
        """
        Normalize whitespace (spaces, tabs, newlines)
        
        Args:
            text: Input text
            
        Returns:
            str: Normalized text
        """
        if not text:
            return ""
        
        # Replace tabs with spaces
        text = text.replace('\t', ' ')
        
        # Normalize multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Normalize multiple newlines (keep max 2)
        text = re.sub(r'\n\n+', '\n\n', text)
        
        return text.strip()
    
    @staticmethod
    def remove_urls(text: str) -> str:
        """Loại bỏ URLs"""
        if not text:
            return ""
        
        # Remove URLs
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
        text = re.sub(r'www\.(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
        
        return text.strip()
    
    @staticmethod
    def remove_email(text: str) -> str:
        """Loại bỏ email addresses"""
        if not text:
            return ""
        
        text = re.sub(r'\S+@\S+', '', text)
        return text.strip()
    
    @staticmethod
    def remove_extra_newlines(text: str, max_newlines: int = 2) -> str:
        """
        Giới hạn số newlines liên tiếp
        
        Args:
            text: Input text
            max_newlines: Số newlines tối đa cho phép
            
        Returns:
            str: Text với newlines được giới hạn
        """
        if not text:
            return ""
        
        pattern = r'\n{' + str(max_newlines + 1) + r',}'
        replacement = '\n' * max_newlines
        
        return re.sub(pattern, replacement, text)
    
    @staticmethod
    def apply_preprocessing(
        text: str,
        clean: bool = True,
        remove_special_chars: bool = False,
        normalize_ws: bool = True,
        remove_urls: bool = True,
        remove_emails: bool = True,
        custom_processors: Optional[List[Callable[[str], str]]] = None
    ) -> str:
        """
        Apply multiple preprocessing steps
        
        Args:
            text: Input text
            clean: Apply basic cleaning
            remove_special_chars: Remove special characters
            normalize_ws: Normalize whitespace
            remove_urls: Remove URLs
            remove_emails: Remove emails
            custom_processors: List of custom processing functions
            
        Returns:
            str: Preprocessed text
        """
        if not text:
            return ""
        
        # Apply built-in processors
        if remove_urls:
            text = DocumentPreprocessor.remove_urls(text)
        
        if remove_emails:
            text = DocumentPreprocessor.remove_email(text)
        
        if remove_special_chars:
            text = DocumentPreprocessor.remove_special_characters(text)
        
        if normalize_ws:
            text = DocumentPreprocessor.normalize_whitespace(text)
        
        if clean:
            text = DocumentPreprocessor.clean_text(text)
        
        # Apply custom processors
        if custom_processors:
            for processor in custom_processors:
                try:
                    text = processor(text)
                except Exception as e:
                    logger.warning(f"Custom processor failed: {e}")
        
        return text


class DocumentStructure:
    """
    Class phân tích và detect cấu trúc document
    """
    
    @staticmethod
    def detect_headers(text: str) -> List[tuple]:
        """
        Detect headers trong document (Markdown style, numbered sections, etc.)
        
        Args:
            text: Document text
            
        Returns:
            List[tuple]: List of (header_text, level, start_pos, end_pos)
        """
        headers = []
        lines = text.split('\n')
        current_pos = 0
        
        for line in lines:
            line_stripped = line.strip()
            
            # Markdown headers (# Header)
            if line_stripped.startswith('#'):
                level = len(line_stripped) - len(line_stripped.lstrip('#'))
                header_text = line_stripped.lstrip('#').strip()
                if header_text:
                    headers.append((
                        header_text,
                        level,
                        current_pos,
                        current_pos + len(line)
                    ))
            
            # Numbered sections (1. Section, 1.1 Subsection)
            elif re.match(r'^\d+(\.\d+)*\.?\s+\w+', line_stripped):
                match = re.match(r'^(\d+(\.\d+)*\.?)\s+(.+)$', line_stripped)
                if match:
                    number = match.group(1)
                    header_text = match.group(3)
                    level = number.count('.') + 1
                    headers.append((
                        header_text,
                        level,
                        current_pos,
                        current_pos + len(line)
                    ))
            
            # ALL CAPS headers
            elif line_stripped and line_stripped.isupper() and len(line_stripped) > 3:
                # Check if it's likely a header (not too long)
                if len(line_stripped) < 100 and not line_stripped.endswith('.'):
                    headers.append((
                        line_stripped,
                        2,  # Assume level 2
                        current_pos,
                        current_pos + len(line)
                    ))
            
            current_pos += len(line) + 1  # +1 for newline
        
        return headers
    
    @staticmethod
    def split_by_headers(text: str, headers: List[tuple]) -> List[tuple]:
        """
        Split document thành sections dựa trên headers
        
        Args:
            text: Document text
            headers: List of headers từ detect_headers()
            
        Returns:
            List[tuple]: List of (section_text, header_text, level, metadata)
        """
        if not headers:
            return [(text, None, 0, {})]
        
        sections = []
        
        for i, (header_text, level, start_pos, end_pos) in enumerate(headers):
            # Find content between this header and next header
            content_start = end_pos
            
            if i < len(headers) - 1:
                content_end = headers[i + 1][2]  # Start of next header
            else:
                content_end = len(text)
            
            section_content = text[content_start:content_end].strip()
            
            metadata = {
                'header': header_text,
                'level': level,
                'start_pos': content_start,
                'end_pos': content_end,
                'header_index': i
            }
            
            sections.append((section_content, header_text, level, metadata))
        
        return sections
    
    @staticmethod
    def split_by_paragraphs(text: str, min_paragraph_length: int = 50) -> List[str]:
        """
        Split text thành paragraphs
        
        Args:
            text: Document text
            min_paragraph_length: Độ dài tối thiểu của paragraph
            
        Returns:
            List[str]: Danh sách paragraphs
        """
        # Split by double newlines
        paragraphs = text.split('\n\n')
        
        # Filter và clean
        cleaned_paragraphs = []
        for para in paragraphs:
            para = para.strip()
            if len(para) >= min_paragraph_length:
                cleaned_paragraphs.append(para)
        
        return cleaned_paragraphs
    
    @staticmethod
    def split_by_sentences(text: str) -> List[str]:
        """
        Split text thành sentences (fallback nếu không có underthesea)
        
        Args:
            text: Document text
            
        Returns:
            List[str]: Danh sách sentences
        """
        try:
            from underthesea import sent_tokenize
            return sent_tokenize(text)
        except ImportError:
            # Fallback regex-based splitting
            sentences = re.split(r'[.!?]+\s+', text)
            return [s.strip() for s in sentences if s.strip()]
    
    @staticmethod
    def analyze_structure(text: str) -> dict:
        """
        Phân tích toàn bộ cấu trúc document
        
        Args:
            text: Document text
            
        Returns:
            dict: Dictionary chứa thông tin cấu trúc
        """
        headers = DocumentStructure.detect_headers(text)
        sections = DocumentStructure.split_by_headers(text, headers)
        paragraphs = DocumentStructure.split_by_paragraphs(text)
        
        return {
            'headers': headers,
            'sections': sections,
            'paragraphs': paragraphs,
            'num_sections': len(sections),
            'num_paragraphs': len(paragraphs),
            'has_structure': len(headers) > 0
        }


class Document:
    """Lớp đại diện cho một tài liệu"""
    def __init__(self, content: str, metadata: Optional[dict] = None):
        self.content = content
        self.metadata = metadata or {}


class DocumentLoader(ABC):
    """Abstract class cho document loader"""
    
    def __init__(
        self, 
        file_path: str,
        enable_preprocessing: bool = True,
        preprocessing_config: Optional[dict] = None
    ):
        """
        Args:
            file_path: Đường dẫn file
            enable_preprocessing: Bật preprocessing
            preprocessing_config: Config cho preprocessing (dict of bool flags)
        """
        self.file_path = Path(file_path)
        
        if not self.file_path.exists():
            raise FileNotFoundError(f"File không tồn tại: {file_path}")
        
        if not self.file_path.is_file():
            raise ValueError(f"Path không phải là file: {file_path}")
        
        self.enable_preprocessing = enable_preprocessing
        self.preprocessing_config = preprocessing_config or {
            'clean': True,
            'remove_special_chars': False,
            'normalize_ws': True,
            'remove_urls': True,
            'remove_emails': True
        }
        
        self.preprocessor = DocumentPreprocessor()
    
    def _preprocess_text(self, text: str) -> str:
        """
        Apply preprocessing to text
        
        Args:
            text: Raw text
            
        Returns:
            str: Preprocessed text
        """
        if not self.enable_preprocessing or not text:
            return text
        
        try:
            return self.preprocessor.apply_preprocessing(
                text,
                **self.preprocessing_config
            )
        except Exception as e:
            logger.warning(f"Preprocessing failed: {e}. Returning original text.")
            return text
    
    @abstractmethod
    def load(self) -> List[Document]:
        """
        Load và trả về danh sách các documents
        
        Returns:
            List[Document]: Danh sách các documents được load
        """
        pass
    
    def load_and_split(
        self, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200,
        splitter=None
    ) -> List[Document]:
        """
        Load và split document thành các phần nhỏ hơn
        
        Args:
            chunk_size: Kích thước mỗi chunk
            chunk_overlap: Overlap giữa các chunks
            splitter: Custom splitter object (optional)
            
        Returns:
            List[Document]: Danh sách các documents đã được split
        """
        try:
            documents = self.load()
            
            if not documents:
                logger.warning(f"No documents loaded from {self.file_path}")
                return []
            
            # Nếu có custom splitter, sử dụng nó
            if splitter:
                return splitter.split_documents(documents)
            
            # Default splitting logic
            split_documents = []
            
            for doc in documents:
                text = doc.content
                
                if not text or not text.strip():
                    continue
                
                # Simple character-based splitting
                for i in range(0, len(text), chunk_size - chunk_overlap):
                    chunk = text[i:i + chunk_size]
                    if chunk.strip():
                        metadata = doc.metadata.copy()
                        metadata['chunk_index'] = i // (chunk_size - chunk_overlap)
                        metadata['chunk_size'] = len(chunk)
                        split_documents.append(Document(content=chunk, metadata=metadata))
            
            return split_documents
            
        except Exception as e:
            logger.error(f"Error in load_and_split: {e}")
            raise
    
    def load_and_split_by_structure(
        self,
        split_by: str = "sections",
        min_chunk_size: int = 100,
        max_chunk_size: Optional[int] = None,
        merge_small_chunks: bool = True
    ) -> List[Document]:
        """
        Load và split document theo cấu trúc (headers, sections, paragraphs)
        
        Args:
            split_by: Cách split ('sections', 'paragraphs', 'headers')
            min_chunk_size: Kích thước tối thiểu của chunk
            max_chunk_size: Kích thước tối đa (optional, sẽ split nếu vượt quá)
            merge_small_chunks: Merge các chunks nhỏ hơn min_chunk_size
            
        Returns:
            List[Document]: Documents đã được split theo cấu trúc
        """
        try:
            documents = self.load()
            
            if not documents:
                logger.warning(f"No documents loaded from {self.file_path}")
                return []
            
            split_documents = []
            
            for doc in documents:
                text = doc.content
                
                if not text or not text.strip():
                    continue
                
                # Analyze structure
                structure = DocumentStructure.analyze_structure(text)
                
                chunks_data = []
                
                if split_by == "sections" and structure['has_structure']:
                    # Split by sections (headers)
                    for section_content, header, level, section_meta in structure['sections']:
                        if len(section_content) >= min_chunk_size:
                            chunks_data.append({
                                'content': section_content,
                                'header': header,
                                'level': level,
                                'type': 'section',
                                **section_meta
                            })
                        elif merge_small_chunks and chunks_data:
                            # Merge với chunk trước
                            chunks_data[-1]['content'] += '\n\n' + section_content
                        elif not merge_small_chunks:
                            chunks_data.append({
                                'content': section_content,
                                'header': header,
                                'level': level,
                                'type': 'section',
                                **section_meta
                            })
                
                elif split_by == "paragraphs":
                    # Split by paragraphs
                    paragraphs = structure['paragraphs']
                    current_chunk = []
                    current_size = 0
                    
                    for para in paragraphs:
                        para_size = len(para)
                        
                        if max_chunk_size and current_size + para_size > max_chunk_size:
                            # Save current chunk
                            if current_chunk:
                                chunks_data.append({
                                    'content': '\n\n'.join(current_chunk),
                                    'type': 'paragraph_group',
                                    'num_paragraphs': len(current_chunk)
                                })
                            current_chunk = [para]
                            current_size = para_size
                        else:
                            current_chunk.append(para)
                            current_size += para_size
                    
                    # Add last chunk
                    if current_chunk:
                        chunks_data.append({
                            'content': '\n\n'.join(current_chunk),
                            'type': 'paragraph_group',
                            'num_paragraphs': len(current_chunk)
                        })
                
                elif split_by == "headers":
                    # Chỉ extract headers và nội dung ngay sau header
                    if structure['has_structure']:
                        for section_content, header, level, section_meta in structure['sections']:
                            # Lấy một phần content sau header
                            preview_length = max_chunk_size if max_chunk_size else 500
                            content_preview = section_content[:preview_length]
                            
                            chunks_data.append({
                                'content': f"{header}\n\n{content_preview}",
                                'header': header,
                                'level': level,
                                'type': 'header_section',
                                **section_meta
                            })
                    else:
                        # Fallback: split by paragraphs
                        chunks_data = [{
                            'content': text,
                            'type': 'full_document'
                        }]
                
                else:
                    # Fallback: return full text
                    chunks_data = [{'content': text, 'type': 'full_document'}]
                
                # Create Document objects
                for i, chunk_data in enumerate(chunks_data):
                    content = chunk_data.pop('content')
                    metadata = doc.metadata.copy()
                    metadata.update({
                        'chunk_index': i,
                        'total_chunks': len(chunks_data),
                        'split_method': split_by,
                        **chunk_data
                    })
                    split_documents.append(Document(content=content, metadata=metadata))
            
            logger.info(f"Split into {len(split_documents)} chunks using '{split_by}' method")
            return split_documents
            
        except Exception as e:
            logger.error(f"Error in load_and_split_by_structure: {e}")
            raise


class PDFLoader(DocumentLoader):
    """
    Loader cho file PDF với error handling và preprocessing
    """
    
    def __init__(
        self,
        file_path: str,
        enable_preprocessing: bool = True,
        preprocessing_config: Optional[dict] = None,
        extract_images: bool = False,
        password: Optional[str] = None
    ):
        """
        Args:
            file_path: Đường dẫn file PDF
            enable_preprocessing: Bật preprocessing
            preprocessing_config: Config cho preprocessing
            extract_images: Extract images từ PDF (not implemented yet)
            password: Password nếu PDF được mã hóa
        """
        if PyPDF2 is None:
            raise ImportError(
                "PyPDF2 không được cài đặt. Cài đặt bằng: pip install PyPDF2"
            )
        
        super().__init__(file_path, enable_preprocessing, preprocessing_config)
        self.extract_images = extract_images
        self.password = password
        
        # Validate PDF file
        if self.file_path.suffix.lower() != '.pdf':
            raise ValueError(f"File không phải PDF: {file_path}")
    
    def load(self) -> List[Document]:
        """
        Load nội dung từ file PDF với error handling
        
        Returns:
            List[Document]: Danh sách documents (mỗi page là 1 document)
        """
        documents = []
        
        try:
            with open(self.file_path, 'rb') as file:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    # Handle encrypted PDFs
                    if pdf_reader.is_encrypted:
                        if self.password:
                            try:
                                pdf_reader.decrypt(self.password)
                                logger.info(f"Successfully decrypted PDF: {self.file_path}")
                            except Exception as e:
                                raise ValueError(f"Failed to decrypt PDF with provided password: {e}")
                        else:
                            raise ValueError(f"PDF is encrypted but no password provided: {self.file_path}")
                    
                    total_pages = len(pdf_reader.pages)
                    logger.info(f"Loading PDF with {total_pages} pages: {self.file_path}")
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            # Extract text
                            text = page.extract_text()
                            
                            if not text or not text.strip():
                                logger.warning(f"Page {page_num + 1} is empty or text extraction failed")
                                continue
                            
                            # Apply preprocessing
                            processed_text = self._preprocess_text(text)
                            
                            if not processed_text or not processed_text.strip():
                                logger.warning(f"Page {page_num + 1} is empty after preprocessing")
                                continue
                            
                            metadata = {
                                'source': str(self.file_path),
                                'page': page_num + 1,
                                'total_pages': total_pages,
                                'file_type': 'pdf'
                            }
                            
                            documents.append(Document(content=processed_text, metadata=metadata))
                            
                        except Exception as e:
                            logger.error(f"Error extracting page {page_num + 1}: {e}")
                            continue
                    
                    if not documents:
                        logger.warning(f"No content extracted from PDF: {self.file_path}")
                    else:
                        logger.info(f"Successfully loaded {len(documents)} pages from PDF")
                    
                except Exception as e:
                    logger.error(f"Error reading PDF file: {e}")
                    raise ValueError(f"Failed to read PDF file {self.file_path}: {e}")
        
        except FileNotFoundError:
            logger.error(f"File not found: {self.file_path}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error loading PDF: {e}")
            raise
        
        return documents


class TXTLoader(DocumentLoader):
    """
    Loader cho file TXT với encoding detection và error handling
    """
    
    def __init__(
        self,
        file_path: str,
        encoding: str = 'utf-8',
        enable_preprocessing: bool = True,
        preprocessing_config: Optional[dict] = None,
        auto_detect_encoding: bool = True
    ):
        """
        Args:
            file_path: Đường dẫn file TXT
            encoding: Encoding của file (mặc định: utf-8)
            enable_preprocessing: Bật preprocessing
            preprocessing_config: Config cho preprocessing
            auto_detect_encoding: Tự động detect encoding nếu utf-8 failed
        """
        super().__init__(file_path, enable_preprocessing, preprocessing_config)
        self.encoding = encoding
        self.auto_detect_encoding = auto_detect_encoding
        
        # Validate TXT file
        if self.file_path.suffix.lower() not in ['.txt', '.text', '.docx']:
            logger.warning(f"File extension is not .txt or .docx: {file_path}")
    
    def _detect_encoding(self) -> str:
        """
        Detect file encoding
        
        Returns:
            str: Detected encoding
        """
        try:
            import chardet
            
            with open(self.file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                detected_encoding = result['encoding']
                confidence = result['confidence']
                
                logger.info(f"Detected encoding: {detected_encoding} (confidence: {confidence})")
                return detected_encoding
                
        except ImportError:
            logger.warning("chardet not installed. Install with: pip install chardet")
            return 'utf-8'
        except Exception as e:
            logger.warning(f"Encoding detection failed: {e}")
            return 'utf-8'
    
    def load(self) -> List[Document]:
        """
        Load nội dung từ file TXT hoặc DOCX với error handling
        
        Returns:
            List[Document]: Danh sách documents (mỗi page là 1 document)
        """
        try:
            # Determine file type and load content
            file_type = self.file_path.suffix.lower()
            pages = []
            used_encoding = None
            
            # Load DOCX file
            if file_type == '.docx':
                if DocxDocument is None:
                    raise ImportError(
                        "python-docx is required to load DOCX files. "
                        "Install with: pip install python-docx"
                    )
                
                doc = DocxDocument(str(self.file_path))
                
                # Group paragraphs by page breaks
                current_page = []
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    
                    # Check for page break
                    if para.runs:
                        for run in para.runs:
                            if '\f' in run.text or '\x0c' in run.text:
                                # Save current page and start new one
                                if current_page:
                                    pages.append('\n'.join(current_page))
                                    current_page = []
                                break
                    
                    if text:
                        current_page.append(text)
                
                # Add last page
                if current_page:
                    pages.append('\n'.join(current_page))
                
                # If no explicit page breaks, split by character count
                if len(pages) == 1 and len(pages[0]) > 3000:
                    pages = self._split_into_pages(pages[0])
                
                file_type_str = 'docx'
                
            # Load TXT file
            else:
                text = None
                used_encoding = self.encoding
                
                # Try with specified encoding
                try:
                    with open(self.file_path, 'r', encoding=self.encoding) as file:
                        text = file.read()
                    logger.info(f"Successfully loaded file with {self.encoding} encoding")
                    
                except UnicodeDecodeError as e:
                    if self.auto_detect_encoding:
                        logger.warning(f"Failed to decode with {self.encoding}: {e}")
                        logger.info("Attempting to auto-detect encoding...")
                        
                        detected_encoding = self._detect_encoding()
                        
                        if detected_encoding and detected_encoding != self.encoding:
                            try:
                                with open(self.file_path, 'r', encoding=detected_encoding) as file:
                                    text = file.read()
                                used_encoding = detected_encoding
                                logger.info(f"Successfully loaded with detected encoding: {detected_encoding}")
                            except Exception as e2:
                                logger.error(f"Failed with detected encoding: {e2}")
                                raise
                        else:
                            raise
                    else:
                        raise
                
                if text is None:
                    raise ValueError(f"Failed to load text from {self.file_path}")
                
                # Check if file is empty
                if not text.strip():
                    logger.warning(f"File is empty: {self.file_path}")
                    return []
                
                # Split text into pages
                pages = self._split_into_pages(text)
                file_type_str = 'txt'
            
            # Check if content is empty
            if not pages:
                logger.warning(f"File is empty: {self.file_path}")
                return []
            
            # Create documents from pages
            documents = []
            for page_num, page_text in enumerate(pages, start=1):
                # Apply preprocessing
                processed_text = self._preprocess_text(page_text)
                
                if not processed_text or not processed_text.strip():
                    continue
                
                metadata = {
                    'source': str(self.file_path),
                    'file_type': file_type_str,
                    'file_size': self.file_path.stat().st_size,
                    'page': page_num,
                    'total_pages': len(pages),
                    'char_count': len(processed_text)
                }
                
                # Add encoding info for TXT files
                if used_encoding:
                    metadata['encoding'] = used_encoding
                
                documents.append(Document(content=processed_text, metadata=metadata))
            
            logger.info(f"Successfully loaded {file_type_str.upper()} file: {len(documents)} pages")
            
            return documents
            
        except FileNotFoundError:
            logger.error(f"File not found: {self.file_path}")
            raise
        except Exception as e:
            logger.error(f"Error loading file: {e}")
            raise
    
    def _split_into_pages(self, text: str, chars_per_page: int = 3000) -> List[str]:
        """
        Split text into pages based on form feed character or character count
        
        Args:
            text: Input text
            chars_per_page: Approximate characters per page
            
        Returns:
            List[str]: List of page texts
        """
        # Check if text contains form feed character (page break)
        if '\f' in text:
            pages = text.split('\f')
            return [page.strip() for page in pages if page.strip()]
        
        # Otherwise split by character count
        pages = []
        current_pos = 0
        while current_pos < len(text):
            end_pos = min(current_pos + chars_per_page, len(text))
            
            # Try to break at sentence or paragraph boundary
            if end_pos < len(text):
                # Look for paragraph break
                para_break = text.rfind('\n\n', current_pos, end_pos)
                if para_break > current_pos:
                    end_pos = para_break
                else:
                    # Look for sentence break
                    sentence_break = max(
                        text.rfind('. ', current_pos, end_pos),
                        text.rfind('! ', current_pos, end_pos),
                        text.rfind('? ', current_pos, end_pos)
                    )
                    if sentence_break > current_pos:
                        end_pos = sentence_break + 1
            
            page_text = text[current_pos:end_pos].strip()
            if page_text:
                pages.append(page_text)
            
            current_pos = end_pos
        
        return pages if pages else [text]


class DirectoryLoader:
    """
    Load multiple documents from a directory
    """
    
    def __init__(
        self,
        directory_path: str,
        glob_pattern: str = "**/*",
        loader_cls_mapping: Optional[dict] = None,
        enable_preprocessing: bool = True,
        show_progress: bool = True
    ):
        """
        Args:
            directory_path: Đường dẫn thư mục
            glob_pattern: Pattern để filter files (e.g., "**/*.pdf")
            loader_cls_mapping: Mapping {extension: LoaderClass}
            enable_preprocessing: Bật preprocessing
            show_progress: Hiển thị progress
        """
        self.directory_path = Path(directory_path)
        
        if not self.directory_path.exists():
            raise FileNotFoundError(f"Directory không tồn tại: {directory_path}")
        
        if not self.directory_path.is_dir():
            raise ValueError(f"Path không phải directory: {directory_path}")
        
        self.glob_pattern = glob_pattern
        self.enable_preprocessing = enable_preprocessing
        self.show_progress = show_progress
        
        # Default loader mapping
        self.loader_cls_mapping = loader_cls_mapping or {
            '.pdf': PDFLoader,
            '.txt': TXTLoader,
            '.text': TXTLoader
        }
    
    def load(self) -> List[Document]:
        """
        Load all documents from directory
        
        Returns:
            List[Document]: Tất cả documents
        """
        all_documents = []
        files = list(self.directory_path.glob(self.glob_pattern))
        
        if not files:
            logger.warning(f"No files found matching pattern: {self.glob_pattern}")
            return []
        
        logger.info(f"Found {len(files)} files to load")
        
        for file_path in files:
            if not file_path.is_file():
                continue
            
            # Get loader class for file extension
            ext = file_path.suffix.lower()
            loader_cls = self.loader_cls_mapping.get(ext)
            
            if not loader_cls:
                logger.warning(f"No loader for extension {ext}: {file_path}")
                continue
            
            try:
                # Create loader and load documents
                loader = loader_cls(
                    str(file_path),
                    enable_preprocessing=self.enable_preprocessing
                )
                docs = loader.load()
                all_documents.extend(docs)
                
                if self.show_progress:
                    logger.info(f"Loaded {len(docs)} documents from {file_path.name}")
                    
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
                continue
        
        logger.info(f"Successfully loaded {len(all_documents)} documents from {len(files)} files")
        return all_documents
